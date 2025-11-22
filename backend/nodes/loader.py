import os
import requests
from bs4 import BeautifulSoup
import docx
from langchain_community.document_loaders import PyPDFLoader, TextLoader, Docx2txtLoader
from langchain_core.documents import Document
from typing import List

def load_document(state):
    """
    Loads a document from the file path specified in the state.
    Supports PDF, TXT, and DOCX.
    """
    print("---LOADING DOCUMENT---")
    
    documents: List[Document] = []

    try:
        # Check if URL is provided
        if state.get("url"):
            url = state["url"]
            print(f"Loading content from URL: {url}")
            try:
                
                response = requests.get(url)
                response.raise_for_status()
                
                soup = BeautifulSoup(response.content, 'html.parser')
                
                # Remove script and style elements
                for script in soup(["script", "style"]):
                    script.decompose()
                    
                text = soup.get_text()
                
                # Clean text
                lines = (line.strip() for line in text.splitlines())
                chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                text = '\n'.join(chunk for chunk in chunks if chunk)
                
                documents = [Document(page_content=text, metadata={"source": url})]
                
            except Exception as e:
                print(f"Error loading URL: {e}")
                raise Exception(f"Failed to load URL: {str(e)}")
                
        elif state.get("file_path"):
            file_path = state["file_path"]
            print(f"Loading document: {file_path}")
            
            if file_path.endswith(".pdf"):
                loader = PyPDFLoader(file_path)
                documents = loader.load()
            elif file_path.endswith(".docx"):
                # Custom docx loader
                doc = docx.Document(file_path)
                text = "\n".join([para.text for para in doc.paragraphs])
                documents = [Document(page_content=text, metadata={"source": file_path})]
            elif file_path.endswith(".txt"):
                loader = TextLoader(file_path)
                documents = loader.load()
            else:
                raise ValueError("Unsupported file format")
        else:
            raise ValueError("No file path or URL provided")
            
        print(f"Loaded {len(documents)} pages/documents.")
        # Merge documents into existing state without discarding other keys
        new_state = dict(state)
        new_state["documents"] = documents
        return new_state
        
    except Exception as e:
        print(f"Error loading document: {e}")
        raise e
